
import arabic_reshaper
from bidi.algorithm import get_display
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import fonts
import streamlit as st
from streamlit_option_menu import option_menu
import tempfile
import os
from PIL import Image
import base64
from io import BytesIO
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
import requests
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import chardet
import pytesseract
import time

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Windows example

# Register Arabic font for PDF generation
try:
    pdfmetrics.registerFont(TTFont('Arabic', 'arial.ttf'))
except:
    st.warning("Arabic font not found. Using default which may not render Arabic properly.")

def format_arabic(text):
    """Format Arabic text for proper display"""
    if not text:
        return text
    try:
        reshaped_text = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped_text)
        return bidi_text
    except:
        return text

# Set page config with dark theme
st.set_page_config(
    page_title="SummarAIze",
    page_icon="circular_logo.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

@st.cache_resource
def load_model():
    model_name = "csebuetnlp/mT5_multilingual_XLSum"
    
    # Explicitly disable fast tokenizer
    tokenizer = AutoTokenizer.from_pretrained(model_name, use_fast=False)
    
    # Load the model
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    
    return tokenizer, model

tokenizer, model = load_model()

# Custom CSS for modern dark-themed styling
def local_css(file_name):
    try:
        with open(file_name) as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    except:
        st.markdown("""
        <style>
        /* Main containers */
        .main {
            background-color: rgba(26, 32, 44, 0.95) !important;
            border-radius: 16px !important;
            padding: 2.5rem !important;
            margin: 2rem 0 !important;
            color: white !important;
            box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3) !important;
            backdrop-filter: blur(2px) !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
        }
        
        /* Summary box with gradient border */
        .summary-box {
            background-color: rgba(45, 55, 72, 0.9) !important;
            border-left: 4px solid;
            border-image: linear-gradient(to bottom, #4b8bff, #6c5ce7) 1;
            padding: 1.8rem !important;
            border-radius: 0 12px 12px 0 !important;
            margin: 1.5rem 0 !important;
            box-shadow: 0 4px 16px rgba(0,0,0,0.4) !important;
            color: white !important;
            transition: all 0.3s ease !important;
        }
        
        .summary-box:hover {
            transform: translateY(-2px);
            box-shadow: 0 8px 24px rgba(0,0,0,0.5) !important;
        }
        
        /* Buttons with gradient and animation */
        .stButton>button {
            border-radius: 12px !important;
            padding: 0.7rem 1.5rem !important;
            transition: all 0.3s ease !important;
            background: linear-gradient(135deg, #4b8bff 0%, #6c5ce7 100%) !important;
            color: white !important;
            border: none !important;
            font-weight: 600 !important;
            box-shadow: 0 4px 12px rgba(75, 139, 255, 0.3) !important;
        }
        
        .stButton>button:hover {
            transform: translateY(-3px) !important;
            box-shadow: 0 8px 24px rgba(75, 139, 255, 0.4) !important;
            background: linear-gradient(135deg, #3a6fd9 0%, #5d4ac9 100%) !important;
        }
        
        /* Text areas with glass effect */
        .stTextArea>div>div>textarea {
            border-radius: 12px !important;
            padding: 1.2rem !important;
            background-color: rgba(45, 55, 72, 0.8) !important;
            color: white !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
            backdrop-filter: blur(8px) !important;
            transition: all 0.3s ease !important;
        }
        
        .stTextArea>div>div>textarea:focus {
            border-color: #4b8bff !important;
            box-shadow: 0 0 0 2px rgba(75, 139, 255, 0.3) !important;
        }
        
        /* Input fields */
        .stSelectbox, .stTextInput, .stSlider {
            margin-bottom: 1.2rem !important;
        }
        
        .stTextInput>div>div>input {
            background-color: rgba(45, 55, 72, 0.8) !important;
            color: white !important;
            border-radius: 12px !important;
            padding: 0.8rem 1rem !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
        }
        
        /* Radio buttons */
        .stRadio>div {
            background-color: rgba(45, 55, 72, 0.8) !important;
            padding: 1.2rem !important;
            border-radius: 12px !important;
            color: white !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
        }
        
        /* Headers */
        .stMarkdown h1, .stMarkdown h2, .stMarkdown h3 {
            color: #ffffff !important;
            font-weight: 700 !important;
        }
        
        /* Dataframes */
        .stDataFrame {
            border-radius: 12px !important;
            box-shadow: 0 4px 16px rgba(0,0,0,0.3) !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
        }
        
        /* Sidebar */
        .sidebar .sidebar-content {
            background-color: rgba(26, 32, 44, 0.95) !important;
            backdrop-filter: blur(8px) !important;
            border-right: 1px solid rgba(255, 255, 255, 0.1) !important;
        }
        
        /* Slider */
        /* ----- تعديلات السلايدرز الجديدة ----- */
        .stSlider {
            margin: 1.5rem 0 2.5rem 0;
        }

        .stSlider .st-ae {
            background: linear-gradient(90deg, #ffffff 0%, #6c5ce7 100%) !important;
            height: 8px !important;
            border-radius: 4px !important;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1) !important;
        }

        .stSlider .st-ae:hover {
            box-shadow: 0 0 12px rgba(75, 139, 255, 0.6) !important;
        }

        .stSlider .st-ag {
            background: rgba(135deg, #4b8bff, #6c5ce7) !important;
            height: 8px !important;
            border-radius: 4px !important;
        }

        .stSlider .st-af {
            color: white !important;
            font-weight: 600 !important;
            margin-bottom: 8px !important;
        }

        .stSlider .st-ah {
            color: rgba(255, 255, 255, 0.8) !important;
            font-size: 0.85rem !important;
            margin-top: 8px !important;
        }

        .stSlider .st-ai {
            color: white !important;
            font-weight: bold !important;
            background: rgba(45, 55, 72, 0.9) !important;
            padding: 4px 8px !important;
            border-radius: 12px !important;
            border: 1px solid rgba(255, 255, 255, 0.1) !important;
            box-shadow: 0 2px 8px rgba(0,0,0,0.2) !important;
        }
        
        /* Progress bar */
        .stProgress > div > div > div > div {
            background: linear-gradient(90deg, #4b8bff, #6c5ce7) !important;
        }
        
        /* Custom scrollbar */
        ::-webkit-scrollbar {
            width: 8px;
            height: 8px;
        }
        
        ::-webkit-scrollbar-track {
            background: rgba(26, 32, 44, 0.8);
        }
        
        ::-webkit-scrollbar-thumb {
            background: linear-gradient(#4b8bff, #6c5ce7);
            border-radius: 10px;
        }
        
        /* Floating animation for logo */
        @keyframes float {
            0% { transform: translateY(0px); }
            50% { transform: translateY(-10px); }
            100% { transform: translateY(0px); }
        }
        
        .floating-logo {
            animation: float 6s ease-in-out infinite;
        }
        
        /* Pulse animation for important buttons */
        @keyframes pulse {
            0% { box-shadow: 0 0 0 0 rgba(75, 139, 255, 0.7); }
            70% { box-shadow: 0 0 0 15px rgba(75, 139, 255, 0); }
            100% { box-shadow: 0 0 0 0 rgba(75, 139, 255, 0); }
        }
        
        .pulse-button {
            animation: pulse 2s infinite;
        }
        
        /* Tooltip styling */
        .tooltip {
            position: relative;
            display: inline-block;
        }
        
        .tooltip .tooltiptext {
            visibility: hidden;
            width: 200px;
            background-color: rgba(45, 55, 72, 0.95);
            color: #fff;
            text-align: center;
            border-radius: 6px;
            padding: 8px;
            position: absolute;
            z-index: 1;
            bottom: 125%;
            left: 50%;
            transform: translateX(-50%);
            opacity: 0;
            transition: opacity 0.3s;
            border: 1px solid rgba(255, 255, 255, 0.1);
            box-shadow: 0 4px 16px rgba(0,0,0,0.3);
            font-size: 14px;
        }
        
        .tooltip:hover .tooltiptext {
            visibility: visible;
            opacity: 1;
        }
        
        /* Feature cards */
        .feature-card {
            background: rgba(45, 55, 72, 0.8);
            border-radius: 12px;
            padding: 1.5rem;
            margin: 1rem 0;
            border: 1px solid rgba(255, 255, 255, 0.1);
            transition: all 0.3s ease;
        }
        
        .feature-card:hover {
            transform: translateY(-5px);
            box-shadow: 0 8px 24px rgba(0,0,0,0.4);
            border-color: rgba(75, 139, 255, 0.3);
        }
        
        .feature-icon {
            font-size: 2.5rem;
            margin-bottom: 1rem;
            background: linear-gradient(135deg, #4b8bff, #6c5ce7);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }
        
        /* Stats counter */
        .stats-container {
            display: flex;
            justify-content: space-around;
            flex-wrap: wrap;
            margin: 2rem 0;
        }
        
        .stat-box {
            text-align: center;
            padding: 1.5rem;
            border-radius: 12px;
            background: rgba(45, 55, 72, 0.8);
            border: 1px solid rgba(255, 255, 255, 0.1);
            min-width: 150px;
            margin: 0.5rem;
            transition: all 0.3s ease;
        }
        
        .stat-box:hover {
            transform: translateY(-5px);
            box-shadow: 0 8px 24px rgba(0,0,0,0.4);
        }
        
        .stat-number {
            font-size: 2.5rem;
            font-weight: 700;
            background: linear-gradient(135deg, #4b8bff, #6c5ce7);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 0.5rem;
        }
        
        .stat-label {
            font-size: 0.9rem;
            color: rgba(255, 255, 255, 0.8);
        }
        </style>
        """, unsafe_allow_html=True)

local_css("style.css")

# Background image from local file with overlay
def set_bg_from_local(image_path):
    if os.path.exists(image_path):
        with open(image_path, "rb") as f:
            image_data = f.read()
        b64_image = base64.b64encode(image_data).decode()
        
        st.markdown(
            f"""
            <style>
            .stApp {{
                background-image: url("data:image/jpeg;base64,{b64_image}");
                background-size: cover;
                background-position: center;
                background-repeat: no-repeat;
                background-attachment: fixed;
            }}
            .main {{
                background-color: rgba(26, 32, 44, 0.95) !important;
                border-radius: 16px !important;
                padding: 2.5rem !important;
                margin: 2rem 0 !important;
                color: white !important;
                box-shadow: 0 8px 32px rgba(0, 0, 0, 0.3) !important;
                backdrop-filter: blur(2px) !important;
                border: 1px solid rgba(255, 255, 255, 0.1) !important;
            }}
            </style>
            """,
            unsafe_allow_html=True
        )

# Set background from local file
set_bg_from_local("SummerAize background black.png")

def image_to_base64(image_path):
    with open(image_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode('utf-8')

# Helper function to create PDF - FIXED VERSION
def text_to_pdf(text, filename):
    """Create PDF file from text with proper Arabic support"""
    try:
        c = canvas.Canvas(filename, pagesize=letter)
        width, height = letter
        margin = 50
        x = margin
        y = height - margin
        line_height = 15
        
        is_arabic = any('\u0600' <= char <= '\u06FF' for char in text)
        
        if is_arabic:
            c.setFont('Arabic', 12)
            text = format_arabic(text)
            lines = []
            for line in text.split('\n'):
                while len(line) > 80:
                    lines.append(line[:80])
                    line = line[80:]
                if line:
                    lines.append(line)
            lines = lines[::-1]
        else:
            c.setFont('Helvetica', 12)
            lines = []
            for line in text.split('\n'):
                while len(line) > 80:
                    lines.append(line[:80])
                    line = line[80:]
                if line:
                    lines.append(line)
        
        for line in lines:
            if y < margin:
                c.showPage()
                y = height - margin
                if is_arabic:
                    c.setFont('Arabic', 12)
                else:
                    c.setFont('Helvetica', 12)
            
            if is_arabic:
                text_width = c.stringWidth(line, 'Arabic', 12)
                c.drawString(width - margin - text_width, y, line)
            else:
                c.drawString(x, y, line)
            
            y -= line_height
        
        c.save()
        return True
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return False

def generate_summary(text, max_length=150, min_length=30):
    inputs = tokenizer.encode("summarize: " + text, return_tensors="pt", max_length=512, truncation=True)
    summary_ids = model.generate(
        inputs,
        max_length=max_length,
        min_length=min_length,
        length_penalty=2.0,
        num_beams=4,
        early_stopping=True
    )
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
    return summary

def process_image_file(image_file):
    image = Image.open(image_file)
    extracted_text = pytesseract.image_to_string(image, lang='eng+ara')  # Supports Arabic & English
    return extracted_text

def process_text_file(file):
    raw_data = file.read()
    result = chardet.detect(raw_data)
    encoding = result['encoding'] if result['confidence'] > 0.7 else 'utf-8'
    file.seek(0)
    text = file.read().decode(encoding)
    return text

def process_pdf_file(file, page_range=None):
    pdf_reader = PdfReader(file)
    page_summaries = []
    full_text = ""
    
    if page_range:
        if '-' in page_range:
            start_page, end_page = map(int, page_range.split('-'))
            pages = range(start_page-1, min(end_page, len(pdf_reader.pages)))
        else:
            page_num = int(page_range) - 1
            pages = [page_num] if 0 <= page_num < len(pdf_reader.pages) else range(len(pdf_reader.pages))
    else:
        pages = range(len(pdf_reader.pages))
    
    for page_num in pages:
        page_text = pdf_reader.pages[page_num].extract_text()
        if page_text.strip():
            page_summary = generate_summary(page_text)
            page_summaries.append({
                "Page": page_num + 1,
                "Original_Text": page_text,
                "Summary": page_summary
            })
            full_text += f"Page {page_num + 1}:\n{page_text}\n\n"
    
    summary_df = pd.DataFrame(page_summaries)
    return full_text, summary_df

def process_word_file(file, page_range=None):
    doc = Document(file)
    page_summaries = []
    full_text = ""
    
    # Get all non-empty paragraphs and track sections
    all_paragraphs = []
    current_section = 1
    for para in doc.paragraphs:
        if para.text.strip():
            # Check for section breaks (approximate)
            if '\\section' in para.style.name.lower() or '\\page' in para.text.lower():
                current_section += 1
            all_paragraphs.append({
                'text': para.text,
                'section': current_section
            })
    
    # Estimate pages based on paragraph count (more accurate than fixed number)
    avg_paragraphs_per_page = max(3, len(all_paragraphs) // 3)  # At least 3 paragraphs per page
    total_pages = max(1, (len(all_paragraphs) // avg_paragraphs_per_page))
    
    # Parse page range
    if page_range:
        if '-' in page_range:
            start_page, end_page = map(int, page_range.split('-'))
            end_page = min(end_page, total_pages)
        else:
            start_page = end_page = int(page_range)
    else:
        start_page = 1
        end_page = total_pages
    
    # Validate range
    if start_page < 1 or end_page > total_pages or start_page > end_page:
        st.error(f"Invalid page range (1-{total_pages} available)")
        return None, None
    
    # Process each "page" in range
    for page_num in range(start_page, end_page + 1):
        start_idx = (page_num - 1) * avg_paragraphs_per_page
        end_idx = min(page_num * avg_paragraphs_per_page, len(all_paragraphs))
        page_paragraphs = all_paragraphs[start_idx:end_idx]
        
        page_text = "\n".join([p['text'] for p in page_paragraphs])
        if page_text.strip():
            try:
                page_summary = generate_summary(page_text)
                page_summaries.append({
                    "Page": page_num,
                    "Original_Text": page_text,
                    "Summary": page_summary
                })
                full_text += f"Page {page_num}:\n{page_text}\n\n"
            except Exception as e:
                st.error(f"Error summarizing page {page_num}: {str(e)}")
                continue
    
    if not page_summaries:
        st.error("No valid content found in selected pages.")
        return None, None
    
    return full_text, pd.DataFrame(page_summaries)

def detect_file_encoding(file):
    raw_data = file.read(10000)
    result = chardet.detect(raw_data)
    file.seek(0)
    return result['encoding'] if result['confidence'] > 0.7 else 'utf-8'

def process_csv_excel(file, file_type, text_column):
    try:
        encoding = detect_file_encoding(file)
        
        if file_type == "csv":
            try:
                df = pd.read_csv(file, encoding=encoding)
            except UnicodeDecodeError:
                for enc in ['utf-8', 'utf-16', 'latin1', 'iso-8859-1', 'cp1252', 'cp1256']:
                    try:
                        file.seek(0)
                        df = pd.read_csv(file, encoding=enc)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    raise ValueError("Could not determine CSV encoding")
        else:
            df = pd.read_excel(file)
        
        if text_column not in df.columns:
            st.error(f"Column '{text_column}' not found in the file.")
            return None
        
        df[text_column] = df[text_column].astype(str).fillna('')
        
        preview_df = df.head(3).copy()
        preview_df['Summary'] = preview_df[text_column].apply(lambda x: generate_summary(str(x)))
        
        df['Summary'] = df[text_column].apply(lambda x: generate_summary(str(x)))
        
        return preview_df, df
    
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

def show_loading_animation():
    """Show a modern loading animation"""
    with st.empty():
        for i in range(3):
            st.markdown("""
            <div style="text-align: center; margin: 2rem 0;">
                <div style="display: inline-block; width: 50px; height: 50px; border: 4px solid rgba(75, 139, 255, 0.3); border-radius: 50%; border-top-color: #4b8bff; animation: spin 1s linear infinite;"></div>
                <p style="margin-top: 1rem; color: white;">Generating your summary...</p>
            </div>
            <style>
                @keyframes spin {
                    0% { transform: rotate(0deg); }
                    100% { transform: rotate(360deg); }
                }
            </style>
            """, unsafe_allow_html=True)
            time.sleep(0.5)

def main():
    if 'summary' not in st.session_state:
        st.session_state.summary = None
    if 'processed_df' not in st.session_state:
        st.session_state.processed_df = None
    if 'source_text' not in st.session_state:
        st.session_state.source_text = None
    if 'page_summaries' not in st.session_state:
        st.session_state.page_summaries = None
    
    with st.sidebar:
        st.markdown("""
        <div style="text-align: center; margin-bottom: 2rem;">
            <h1 style="display: inline; background: linear-gradient(135deg, #4b8bff, #6c5ce7); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 800;">⚙️ Settings</h1>
        </div>
        """, unsafe_allow_html=True)
        
        st.subheader("Summary Settings")
        max_length = st.slider("Maximum summary length", 50, 300, 150, help="Adjust the maximum length of your summary")
        min_length = st.slider("Minimum summary length", 10, 100, 30, help="Adjust the minimum length of your summary")
        
        st.subheader("Language Support")
        languages = {
            "Arabic": "ar",
            "English": "en",
            "Spanish": "es",
            "French": "fr",
            "German": "de",
            "Russian": "ru",
            "Chinese": "zh",
            "Hindi": "hi"
        }
        selected_lang = st.selectbox("Input Language", list(languages.keys()), help="Select the language of your input text")
        
        st.markdown("---")
        st.markdown("""
        <div style="text-align: center; margin-top: 2rem;">
            <p style="font-size: 0.9rem; color: rgba(255, 255, 255, 0.7);">Powered by</p>
            <a href="https://huggingface.co/csebuetnlp/mT5_multilingual_XLSum" target="_blank" style="color: #4b8bff; text-decoration: none; font-weight: 600;">CSEBUET NLP's mT5 model</a>
        </div>
        """, unsafe_allow_html=True)

    # Hero section with floating logo
    col1, col2 = st.columns([1, 3])
    with col1:
        try:
            logo_base64 = image_to_base64("circular_logo.png")
            st.markdown(f"""
            <div style="text-align: center;">
                <div class="floating-logo">
                    <img src="data:image/png;base64,{logo_base64}" width="180">
                </div>
            </div>
            """, unsafe_allow_html=True)
        except Exception as e:
            st.warning(f"Could not load logo: {str(e)}")
            st.markdown("""
            <div style="text-align: center;">
                <div class="floating-logo">
                    <svg width="100" height="100" viewBox="0 0 100 100" fill="none" xmlns="http://www.w3.org/2000/svg">
                        <circle cx="50" cy="50" r="45" fill="none" stroke="url(#gradient)" stroke-width="8"/>
                        <path d="M30 40 L50 70 L70 40" stroke="url(#gradient)" stroke-width="6" stroke-linecap="round"/>
                        <defs>
                            <linearGradient id="gradient" x1="0%" y1="0%" x2="100%" y2="100%">
                                <stop offset="0%" stop-color="#4b8bff"/>
                                <stop offset="100%" stop-color="#6c5ce7"/>
                            </linearGradient>
                        </defs>
                    </svg>
                </div>
            </div>
            """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <h1 style="font-size: 3rem; margin-bottom: 0.5rem; background: linear-gradient(135deg, #4b8bff, #6c5ce7); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 800;">SummarAIze</h1>
        <p style="font-size: 1.2rem; color: rgba(255, 255, 255, 0.8); margin-bottom: 1.5rem;">AI-Powered Multilingual Article Summarization</p>
        """, unsafe_allow_html=True)
    
    # Stats counter
    st.markdown("""
    <div class="stats-container">
        <div class="stat-box">
            <div class="stat-number">7+</div>
            <div class="stat-label">Languages</div>
        </div>
        <div class="stat-box">
            <div class="stat-number">6</div>
            <div class="stat-label">File Types</div>
        </div>
        <div class="stat-box">
            <div class="stat-number">∞</div>
            <div class="stat-label">Possibilities</div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Feature cards
    with st.expander("✨ Key Features", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown("""
            <div class="feature-card">
                <div class="feature-icon">🌐</div>
                <h3>Multilingual</h3>
                <p>Supports multiple languages including Arabic, English, Spanish, and more.</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("""
            <div class="feature-card">
                <div class="feature-icon">📄</div>
                <h3>Multiple Formats</h3>
                <p>Works with PDFs, Word docs, text files, images, and spreadsheets.</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col3:
            st.markdown("""
            <div class="feature-card">
                <div class="feature-icon">⚡</div>
                <h3>Fast & Accurate</h3>
                <p>Powered by state-of-the-art AI models for high-quality summaries.</p>
            </div>
            """, unsafe_allow_html=True)
    
    selected = option_menu(
        menu_title=None,
        options=["Text Input", "File Upload"],
        icons=["pencil-square", "upload"],
        menu_icon="cast",
        default_index=0,
        orientation="horizontal",
        styles={
            "container": {"padding": "0!important", "background-color": "#1a202c", "border-radius": "12px"},
            "icon": {"color": "orange", "font-size": "18px"},
            "nav-link": {
                "font-size": "16px",
                "text-align": "center",
                "margin": "0px",
                "--hover-color": "#2d3748",
                "color": "white",
                "border-radius": "12px",
                "padding": "12px"
            },
            "nav-link-selected": {
                "background": "linear-gradient(135deg, #4b8bff, #6c5ce7)",
                "color": "white",
                "font-weight": "600"
            },
        }
    )
    
    if selected == "Text Input":
        with st.container():
            st.subheader("Paste Your Article")
            article_text = st.text_area(
                "Enter the article you want to summarize (supports multiple languages):",
                height=300,
                placeholder="Paste your article here...",
                label_visibility="collapsed"
            )
            
            col1, col2 = st.columns([1, 3])
            with col1:
                summarize_btn = st.button(
                    "✨ Generate Summary",
                    use_container_width=True,
                    type="primary",
                    key="text_summarize"
                )
            
            if summarize_btn and article_text:
                with st.spinner("Generating high-quality summary..."):
                    show_loading_animation()
                    st.session_state.source_text = article_text
                    st.session_state.summary = generate_summary(article_text, max_length, min_length)
                    st.session_state.processed_df = None
                    st.session_state.page_summaries = None
                    st.rerun()
            
    else:
        with st.container():
            st.subheader("Upload Your Document")
            
            file_type = st.radio(
                "Select file type:",
                ["PDF", "Word", "Text", "CSV", "Excel", "Image"],
                horizontal=True
            )

            allowed_types = {
                "PDF": ["pdf"],
                "Word": ["docx"],
                "Text": ["txt"],
                "CSV": ["csv"],
                "Excel": ["xlsx"],
                "Image": ["png", "jpg", "jpeg"]
            }

            uploaded_file = st.file_uploader(
                f"Upload {file_type} file",
                type=allowed_types[file_type],
                help=f"Upload a {file_type} file to summarize"
            )

            
            if uploaded_file:
                if file_type in ["PDF", "Word"]:
                    page_range = st.text_input(
                        "Page range (e.g., '1-3' for pages 1 to 3, or '2' for just page 2):",
                        placeholder="1-3 or 2",
                        help="Specify which pages to summarize"
                    )
                
                if file_type in ["CSV", "Excel"]:
                    text_column = st.text_input(
                        "Enter the column name containing articles:",
                        placeholder="e.g., 'article' or 'text'",
                        help="Column containing text to summarize"
                    )
                
                process_btn = st.button(
                    "🚀 Process File",
                    use_container_width=True,
                    type="primary",
                    key="file_process"
                )
                
                if process_btn:
                    try:
                        with st.spinner(f"Processing {file_type} file..."):
                            show_loading_animation()
                            if file_type == "Text":
                                text = process_text_file(uploaded_file)
                                st.session_state.source_text = text
                                st.session_state.summary = generate_summary(text, max_length, min_length)
                                st.session_state.processed_df = None
                                st.session_state.page_summaries = None
                                st.experimental_rerun()
                            
                            elif file_type == "Image":
                                text = process_image_file(uploaded_file)
                                st.subheader("Extracted Text Preview")
                                st.text_area("Text extracted from image:", value=text, height=200)
                                if text.strip():
                                    st.session_state.source_text = text
                                    st.session_state.summary = generate_summary(text, max_length, min_length)
                                    st.session_state.processed_df = None
                                    st.session_state.page_summaries = None
                                    st.experimental_rerun()
                                else:
                                    st.warning("No text could be extracted from the image.")
                                                         
                            elif file_type == "PDF":
                                if not page_range:
                                    st.warning("Please specify a page range for PDF documents")
                                else:
                                    result = process_pdf_file(uploaded_file, page_range)
                                    if result:
                                        full_text, summary_df = result
                                        st.session_state.source_text = full_text
                                        st.session_state.page_summaries = summary_df
                                        st.session_state.summary = None
                                        st.session_state.processed_df = None
                                        st.experimental_rerun()
                            
                            elif file_type == "Word":
                                if not page_range:
                                    st.warning("Please specify a page range for Word documents")
                                else:
                                    result = process_word_file(uploaded_file, page_range)
                                    if result:
                                        full_text, summary_df = result
                                        st.session_state.source_text = full_text
                                        st.session_state.page_summaries = summary_df
                                        st.session_state.summary = None
                                        st.session_state.processed_df = None
                                        st.experimental_rerun()
                            
                            elif file_type in ["CSV", "Excel"]:
                                if not text_column:
                                    st.error("Please specify the column name containing articles")
                                else:
                                    result = process_csv_excel(
                                        uploaded_file,
                                        file_type.lower(),
                                        text_column
                                    )
                                    if result:
                                        preview_df, full_df = result
                                        st.session_state.processed_df = full_df
                                        st.session_state.summary = None
                                        st.session_state.source_text = None
                                        st.session_state.page_summaries = None
                                        st.experimental_rerun()
                    
                    except Exception as e:
                        st.error(f"An error occurred: {str(e)}")
    
    # Display summary and download options (persistent until new summary)
    if st.session_state.summary:
        st.subheader("Generated Summary")
        st.markdown(f'<div class="summary-box">{st.session_state.summary}</div>', unsafe_allow_html=True)
        
        # Create a DataFrame for all download formats
        summary_df = pd.DataFrame({
            "Original_Text": [st.session_state.source_text],
            "Summary": [st.session_state.summary]
        })
        
        # Download options for all formats
        st.markdown("---")
        st.subheader("Download Options")
        
        # Row 1: TXT, Word, PDF
        col1, col2, col3 = st.columns(3)
        with col1:
            # TXT download
            st.download_button(
                "📄 Download as TXT",
                data=st.session_state.summary,
                file_name="summary.txt",
                mime="text/plain",
                key="txt_download"
            )
        with col2:
            # Word download
            doc = Document()
            doc.add_paragraph(st.session_state.summary)
            bio = BytesIO()
            doc.save(bio)
            st.download_button(
                "📝 Download as Word",
                data=bio.getvalue(),
                file_name="summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="word_download"
            )
        with col3:
            # PDF download - using a temporary file with proper cleanup
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_path = tmp_file.name
                
                if text_to_pdf(st.session_state.summary, tmp_path):
                    with open(tmp_path, "rb") as f:
                        pdf_data = f.read()
                    st.download_button(
                        "📑 Download as PDF",
                        data=pdf_data,
                        file_name="summary.pdf",
                        mime="application/pdf",
                        key="pdf_download"
                    )
            finally:
                if os.path.exists(tmp_path):
                    try:
                        os.unlink(tmp_path)
                    except PermissionError:
                        pass  # Skip if file is still in use
        
        # Row 2: CSV and Excel
        col4, col5 = st.columns(2)
        with col4:
            # CSV download with UTF-8 encoding for Arabic
            csv_data = summary_df.to_csv(index=False).encode('utf-8-sig')  # Note utf-8-sig for Excel compatibility
            st.download_button(
                "📊 Download as CSV",
                data=csv_data,
                file_name="summary.csv",
                mime="text/csv",
                key="csv_download"
            )
        with col5:
            # Excel download
            excel_data = BytesIO()
            with pd.ExcelWriter(excel_data, engine='xlsxwriter') as writer:
                summary_df.to_excel(writer, index=False)
            st.download_button(
                "📈 Download as Excel",
                data=excel_data.getvalue(),
                file_name="summary.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="excel_download"
            )
    
    elif st.session_state.page_summaries is not None:
        # Display page summaries and download options
        st.subheader("Page Summaries")
        st.dataframe(st.session_state.page_summaries[['Page', 'Summary']])
        
        # Download options for all page summaries
        st.markdown("---")
        st.subheader("Download Options")
        
        # Full document download
        col1, col2, col3 = st.columns(3)
        with col1:
            # TXT download of all summaries
            all_summaries = "\n\n".join(
                [f"Page {row['Page']}:\n{row['Summary']}" 
                 for _, row in st.session_state.page_summaries.iterrows()]
            )
            st.download_button(
                "📄 All Summaries as TXT",
                data=all_summaries,
                file_name="all_summaries.txt",
                mime="text/plain",
                key="all_txt_download"
            )
        with col2:
            # Word download of all summaries
            doc = Document()
            for _, row in st.session_state.page_summaries.iterrows():
                doc.add_heading(f"Page {row['Page']}", level=2)
                doc.add_paragraph(row['Summary'])
                doc.add_paragraph()
            bio = BytesIO()
            doc.save(bio)
            st.download_button(
                "📝 All Summaries as Word",
                data=bio.getvalue(),
                file_name="all_summaries.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="all_word_download"
            )
        with col3:
            # PDF download of all summaries
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_path = tmp_file.name
                
                all_summaries_text = "\n".join(
                    [f"Page {row['Page']}:\n{row['Summary']}\n\n" 
                     for _, row in st.session_state.page_summaries.iterrows()]
                )
                if text_to_pdf(all_summaries_text, tmp_path):
                    with open(tmp_path, "rb") as f:
                        pdf_data = f.read()
                    st.download_button(
                        "📑 All Summaries as PDF",
                        data=pdf_data,
                        file_name="all_summaries.pdf",
                        mime="application/pdf",
                        key="all_pdf_download"
                    )
            finally:
                if os.path.exists(tmp_path):
                    try:
                        os.unlink(tmp_path)
                    except PermissionError:
                        pass  # Skip if file is still in use
        
        # Full data download (CSV and Excel)
        col4, col5 = st.columns(2)
        with col4:
            # CSV download
            st.download_button(
                "📊 Full Data as CSV",
                data=st.session_state.page_summaries.to_csv(index=False).encode('utf-8-sig'),
                file_name="page_summaries.csv",
                mime="text/csv",
                key="full_csv_download"
            )
        with col5:
            # Excel download
            excel_data = BytesIO()
            with pd.ExcelWriter(excel_data, engine='xlsxwriter') as writer:
                st.session_state.page_summaries.to_excel(writer, index=False)
            st.download_button(
                "📈 Full Data as Excel",
                data=excel_data.getvalue(),
                file_name="page_summaries.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="full_excel_download"
            )
    
    elif st.session_state.processed_df is not None:
        # Download options for processed DataFrame
        st.markdown("---")
        st.subheader("Download Full Results")
        
        col1, col2 = st.columns(2)
        with col1:
            # CSV download
            st.download_button(
                "📊 Download as CSV",
                data=st.session_state.processed_df.to_csv(index=False).encode('utf-8-sig'),
                file_name="summarized_results.csv",
                mime="text/csv",
                key="results_csv_download"
            )
        with col2:
            # Excel download
            excel_data = BytesIO()
            with pd.ExcelWriter(excel_data, engine='xlsxwriter') as writer:
                st.session_state.processed_df.to_excel(writer, index=False)
            st.download_button(
                "📈 Download as Excel",
                data=excel_data.getvalue(),
                file_name="summarized_results.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="results_excel_download"
            )
        
        # Also provide options to download sample as other formats
        st.subheader("Download Sample as Other Formats")
        sample_df = st.session_state.processed_df.head(3)
        
        col1, col2, col3 = st.columns(3)
        with col1:
            # TXT download
            sample_text = "\n\n".join([f"Original:\n{row[0]}\n\nSummary:\n{row[1]}" for row in sample_df.values])
            st.download_button(
                "📄 Sample as TXT",
                data=sample_text,
                file_name="sample_summaries.txt",
                mime="text/plain",
                key="sample_txt_download"
            )
        with col2:
            # Word download
            doc = Document()
            for _, row in sample_df.iterrows():
                doc.add_paragraph("Original:")
                doc.add_paragraph(row[0])
                doc.add_paragraph("Summary:")
                doc.add_paragraph(row[1])
                doc.add_paragraph("---")
            bio = BytesIO()
            doc.save(bio)
            st.download_button(
                "📝 Sample as Word",
                data=bio.getvalue(),
                file_name="sample_summaries.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="sample_word_download"
            )
        with col3:
            # PDF download with proper file handling
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                    tmp_path = tmp_file.name
                
                sample_text = "\n".join([f"Original:\n{row[0]}\n\nSummary:\n{row[1]}\n\n" for row in sample_df.values])
                if text_to_pdf(sample_text, tmp_path):
                    with open(tmp_path, "rb") as f:
                        pdf_data = f.read()
                    st.download_button(
                        "📑 Sample as PDF",
                        data=pdf_data,
                        file_name="sample_summaries.pdf",
                        mime="application/pdf",
                        key="sample_pdf_download"
                    )
            finally:
                if os.path.exists(tmp_path):
                    try:
                        os.unlink(tmp_path)
                    except PermissionError:
                        pass  # Skip if file is still in use

if __name__ == "__main__":
    main()